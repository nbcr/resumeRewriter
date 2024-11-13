import docx
from extract_resume_info import extract_info_from_resume

def generate_resume(job_data, output_dir="output", resume_info=None):
    """Generate a resume document for each job posting."""
    for job in job_data:
        doc = docx.Document()
        doc.add_heading(resume_info['name'], level=1)
        doc.add_paragraph(f"Email: {resume_info['email']}")
        doc.add_paragraph(f"Phone: {resume_info['phone']}")
        doc.add_paragraph(f"Location: {resume_info.get('address', 'N/A')}")

        # Add skills and experience from resume_info
        doc.add_heading('Skills', level=2)
        for skill in resume_info['skills']:
            doc.add_paragraph(skill, style='List Bullet')

        doc.add_heading('Experience', level=2)
        for experience in resume_info['experience']:
            doc.add_paragraph(experience, style='List Bullet')

        # Add job-specific information
        doc.add_heading(job['title'], level=2)
        doc.add_paragraph(f"Company: {job['company']}")
        doc.add_paragraph(f"Location: {job['location']}")
        doc.add_paragraph(f"Salary: {job['salary']}")
        doc.add_paragraph("Job Description:")
        doc.add_paragraph(job.get('description', "Description not provided"))

        # Save the document
        file_name = f"{job['company'].replace(' ', '_')}_{job['title'].replace(' ', '_')}.docx"
        doc.save(f"{output_dir}/{file_name}")
        print(f"Saved resume to {output_dir}/{file_name}")


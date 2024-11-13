import os
import time
import requests
import openai
import docx
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options

# Set up OpenAI API
openai.api_key = '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'

# URL for job postings in North Bay, ON
job_bank_url = 'https://www.jobbank.gc.ca/jobsearch/jobsearch?pst=P1B1A1&d=5&fcan=1&fjsf=0&sort=M&fjap=0&fsrc=16'

# Set up Selenium WebDriver (requires chromedriver to be in PATH or specified path to chromedriver executable)
chrome_options = Options()
chrome_options.add_argument("--headless")  # Run in headless mode
service = Service("C:/Apps/chromedriver/chromedriver.exe")  # Update path if needed
driver = webdriver.Chrome(service=service, options=chrome_options)

from selenium.common.exceptions import NoSuchElementException

def fetch_job_postings(url):
    driver.get(url)
    time.sleep(2)
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    
    jobs = []
    job_links = soup.find_all('a', class_='resultJobItem')
    
    job_count = 0  # Initialize counter for limiting the number of jobs

    for link in job_links:
        if job_count >= 3:  # Stop after processing 3 jobs
            break
        
        job_url = 'https://www.jobbank.gc.ca' + link['href']
        job_title = link.get_text(strip=True)
        
        # Go to job page to reveal contact information
        driver.get(job_url)
        time.sleep(2)
        
        try:
            # Locate and click the "Show how to apply" button by its unique ID
            apply_button = driver.find_element(By.ID, "applynowbutton")
            apply_button.click()
            time.sleep(2)
            
            # Re-scrape the page for email and other contact information
            job_soup = BeautifulSoup(driver.page_source, 'html.parser')
            apply_section = job_soup.find('div', class_='how-to-apply')
            email = apply_section.find('a', href=True).get_text(strip=True) if apply_section else None

            # Find the job description, if available
            job_description_section = job_soup.find('div', class_='job-posting-detail')
            if job_description_section:
                job_description = job_description_section.get_text(strip=True)
            else:
                print(f"Warning: 'job-posting-detail' not found for {job_title} at {job_url}")
                job_description = "Job description not available"
            
            jobs.append({
                'title': job_title, 
                'url': job_url, 
                'email': email, 
                'description': job_description
            })
        
        except NoSuchElementException:
            print(f"Could not find 'Show how to apply' button for job: {job_title}")
            jobs.append({'title': job_title, 'url': job_url, 'email': None, 'description': "Job description not available"})
        
        except Exception as e:
            print(f"An error occurred while retrieving contact info for {job_title}: {e}")
            jobs.append({'title': job_title, 'url': job_url, 'email': None, 'description': "Job description not available"})
        
        job_count += 1  # Increment the job count

    return jobs

def generate_application(master_resume, job_description, company_info):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a resume tailoring assistant."},
            {"role": "user", "content": f"""
                Here is my master resume:

                {master_resume}

                Based on the job description below, create a tailored resume and cover letter for the company, including the company name, address, and email if available.

                Job Description: {job_description}
                Company Info: {company_info}

                The cover letter should:
                - Address the company and hiring manager, if known.
                - Mention the job title and your interest in the role.
                - Highlight relevant skills and experiences.
                - Fit within a one-page format.
            """}
        ],
        max_tokens=2000
    )
    return response['choices'][0]['message']['content'].strip()

# Path to the master resume
master_resume_path = 'C:/Users/Yot/Downloads/Work/Master Resume.docx'
def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

master_resume = read_docx(master_resume_path)

# Fetch job postings
job_postings = fetch_job_postings(job_bank_url)

# Directory to save tailored resumes and cover letters
output_dir = 'C:/Users/Yot/Downloads/Work/Applications'
os.makedirs(output_dir, exist_ok=True)

# Process each job posting
for job in job_postings:
    job_title = job['title']
    job_url = job['url']
    email = job.get('email', 'N/A')
    
    # Fetch job description
    job_response = requests.get(job_url)
    job_soup = BeautifulSoup(job_response.text, 'html.parser')
    job_description = job_soup.find('div', class_='job-posting-detail').get_text(strip=True)
    
    # Retrieve additional company info with GPT API
    company_info = get_company_info(job_title, email)
    
    # Generate tailored resume and cover letter
    application_content = generate_application(master_resume, job_description, company_info)
    
    # Save to a .docx file
    output_path = os.path.join(output_dir, f'Application_{job_title}.docx')
    doc = docx.Document()
    doc.add_paragraph(application_content)
    doc.save(output_path)
    
    print(f'Application saved for job: {job_title}')

# Close the WebDriver
driver.quit()
